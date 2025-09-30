import pandas as pd
import time
from openai import OpenAI
from docx import Document

# 🔑 OpenAI API 키 입력
client = OpenAI(api_key="")

# CSV 파일 불러오기
file_paths = ["225.csv"]
all_data = pd.concat([pd.read_csv(path) for path in file_paths], ignore_index=True)

# question_id에서 숫자만 추출
all_data['question_num'] = all_data['question_id'].str.extract('(\d+)').astype(int)

def analyze_section(qa_pairs):
    """
    섹션 단위로 분석 요약 (GPT가 섹션 주제를 스스로 도출 + 보고서 4요소 포함)
    """
    answers = [a for _, a in qa_pairs if isinstance(a, str)]
    answer_text = "\n".join(answers)

    formatted = "\n".join([
        f"Q: {q}\nA: {a}" for q, a in qa_pairs
    ])

    prompt = f"""
당신은 인터뷰 데이터를 분석하는 전문가입니다.  

다음은 하나의 섹션에 속하는 질문과 답변 모음입니다.  
이 데이터를 분석하여 전문적인 보고서 단락을 작성하세요.  

[요구사항]
1. **섹션 제목**: 질문들이 다루는 주제를 스스로 정할 것
2. **요약 수치**: 응답에서 나타난 주요 패턴과 키워드를 비율/빈도 중심으로 제시 (예: "약 40%가 가격을 언급")
3. **대표 인용문**: 응답자의 답변 중 특징적이거나 인상적인 표현을 1~2개 짧게 인용
4. **긍/부정 요약**: 응답자들이 해당 주제에 대해 긍정적으로 보는 점과 부정적으로 보는 점을 균형 있게 정리
5. **시사점/인사이트**: 응답 결과가 서비스 개선이나 전략 수립에 어떤 의미를 가지는지 제안  

[데이터: 질문+답변]
{formatted}

[데이터: 답변 모음]
{answer_text}

[출력 형식 예시]
- 섹션 제목
- 요약 단락 (5~8문장, 위 요구사항을 포함)
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("Error:", e)
        return None


# Word 문서 생성
doc = Document()
doc.add_heading("인터뷰 분석 보고서", level=0)

# session별 그룹핑 (session 컬럼 있다고 가정)
if "session" not in all_data.columns:
    raise ValueError("CSV에 'session' 컬럼이 필요합니다.")

for session_id, group in all_data.groupby("session"):
    qa_pairs = list(zip(group['question_text'], group['answer']))

    if not qa_pairs:
        continue

    # GPT 분석
    result = analyze_section(qa_pairs)
    if result:
        doc.add_heading(f"세션 {session_id}", level=1)
        doc.add_paragraph(result)

    time.sleep(2)  # 안전한 API 호출 간격

# Word 파일 저장
output_path = "interview_session_report.docx"
doc.save(output_path)

print(f"✅ 분석 완료! Word 보고서가 '{output_path}' 파일로 생성되었습니다.")
